import io
import os
import re
import pandas as pd
import streamlit as st

try:
    import pypdf
except ImportError:
    pypdf = None

try:
    import docx
except ImportError:
    docx = None

try:
    import openpyxl
except ImportError:
    openpyxl = None


LABEL_KEYWORDS = {
    "Instructions to Bidders": [
        "instructions to bidders", "itb", "rfq", "rfp",
        "submission deadline", "how to bid", "bidding procedures",
    ],
    "Technical Specifications": [
        "technical specification", "scope of work", "deliverable",
        "shall", "must", "performance requirement", "technical criteria",
    ],
    "Financial Template": [
        "bill of quantities", "boq", "unit price", "total price",
        "rate", "financial offer", "price schedule",
    ],
    "Legal / Eligibility Forms": [
        "eligibility", "declaration", "certificate", "registration",
        "authorized signatory", "power of attorney",
    ],
    "Annexes / Returnables": [
        "annex", "appendix", "form", "returnable", "attachment", "schedule",
    ],
}


class DocumentIngestor:
    def load_folder(self, folder_path: str) -> dict:
        docs = {}
        if not os.path.isdir(folder_path):
            if "errors" not in st.session_state:
                st.session_state["errors"] = []
            st.session_state["errors"].append(
                {"file": folder_path, "error": "Folder path does not exist or is not a directory"}
            )
            return docs

        supported = (".pdf", ".docx", ".xlsx", ".csv")
        for root, _, files in os.walk(folder_path):
            for fname in files:
                ext = os.path.splitext(fname)[1].lower()
                if ext not in supported:
                    continue
                fpath = os.path.join(root, fname)
                try:
                    text, pages = self._extract(fpath, ext)
                    docs[fname] = {
                        "text": text,
                        "pages": pages,
                        "file_type": ext.lstrip("."),
                        "char_count": len(text),
                        "label": "Unknown",
                        "confidence": 0.0,
                        "path": fpath,
                    }
                except Exception as e:
                    if "errors" not in st.session_state:
                        st.session_state["errors"] = []
                    st.session_state["errors"].append({"file": fname, "error": str(e)})
        return docs

    def load_files(self, uploaded_files: list) -> dict:
        docs = {}
        if "errors" not in st.session_state:
            st.session_state["errors"] = []
        for uf in uploaded_files:
            fname = uf.name
            ext = os.path.splitext(fname)[1].lower()
            supported = (".pdf", ".docx", ".xlsx", ".csv")
            if ext not in supported:
                continue
            try:
                raw = uf.read()
                data = io.BytesIO(raw)
                text, pages = self._extract_fileobj(data, ext)
                docs[fname] = {
                    "text": text,
                    "pages": pages,
                    "file_type": ext.lstrip("."),
                    "char_count": len(text),
                    "label": "Unknown",
                    "confidence": 0.0,
                    "path": fname,
                    "raw_bytes": raw,
                }
            except Exception as e:
                st.session_state["errors"].append({"file": fname, "error": str(e)})
        return docs

    def _extract(self, fpath: str, ext: str):
        if ext == ".pdf":
            return self._extract_pdf(fpath)
        elif ext == ".docx":
            return self._extract_docx(fpath)
        elif ext == ".xlsx":
            return self._extract_xlsx(fpath)
        elif ext == ".csv":
            return self._extract_csv(fpath)
        raise ValueError(f"Unsupported file type: {ext}")

    def _extract_fileobj(self, fileobj: io.BytesIO, ext: str):
        if ext == ".pdf":
            return self._extract_pdf(fileobj)
        elif ext == ".docx":
            return self._extract_docx(fileobj)
        elif ext == ".xlsx":
            return self._extract_xlsx(fileobj)
        elif ext == ".csv":
            return self._extract_csv(fileobj)
        raise ValueError(f"Unsupported file type: {ext}")

    def _extract_pdf(self, source):
        reader = pypdf.PdfReader(source)
        pages = []
        for i, page in enumerate(reader.pages):
            try:
                text = page.extract_text() or ""
            except Exception:
                text = ""
            pages.append(text)
        joined = "\n".join(
            f"\n--- Page {i + 1} ---\n{p}" for i, p in enumerate(pages)
        )
        return joined, pages

    def _extract_docx(self, source):
        document = docx.Document(source)
        parts = []
        for para in document.paragraphs:
            if not para.text.strip():
                continue
            if para.style.name.startswith("Heading"):
                parts.append(f"## {para.text}")
            else:
                parts.append(para.text)
        for table in document.tables:
            for row in table.rows:
                for cell in row.cells:
                    if cell.text.strip():
                        parts.append(cell.text.strip())
        joined = "\n".join(parts)
        return joined, [joined]

    def _extract_xlsx(self, source):
        wb = openpyxl.load_workbook(source, data_only=True)
        sheets = []
        sheet_texts = []
        for name in wb.sheetnames:
            ws = wb[name]
            rows = []
            for row in ws.iter_rows(values_only=True):
                rows.append([str(c) if c is not None else "" for c in row])
            if not rows:
                continue
            try:
                df = pd.DataFrame(rows[1:], columns=rows[0])
                md = df.to_markdown(index=False)
            except Exception:
                md = "\n".join(["\t".join(r) for r in rows])
            sheet_text = f"\n--- Sheet: {name} ---\n{md}"
            sheets.append(sheet_text)
            sheet_texts.append(sheet_text)
        joined = "\n".join(sheets)
        return joined, sheet_texts

    def _extract_csv(self, source):
        df = pd.read_csv(source, encoding="utf-8", on_bad_lines="skip")
        md = df.to_markdown(index=False)
        return md, [md]

    def classify_documents(self, docs: dict) -> dict:
        for fname, meta in docs.items():
            text = meta["text"]
            snippet = (text[:3000] + text[-1000:]).lower()
            best_label = "Unknown"
            best_score = 0.0
            for label, keywords in LABEL_KEYWORDS.items():
                matched = sum(1 for kw in keywords if kw in snippet)
                score = matched / len(keywords)
                if score > best_score:
                    best_score = score
                    best_label = label
            docs[fname]["label"] = best_label
            docs[fname]["confidence"] = round(best_score, 3)
        return docs
